# app/services/document_processors.py
import io
import logging
from typing import Dict, Any, Optional
from fastapi import UploadFile, HTTPException
from PIL import Image
import pytesseract
import openpyxl
from docx import Document as DocxDocument
import PyPDF2

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Base class for document processors"""
    
    @staticmethod
    def get_processor(file: UploadFile) -> 'DocumentProcessor':
        """Factory method to get appropriate processor based on file type"""
        filename = file.filename.lower()
        
        if filename.endswith(('.jpg', '.jpeg', '.png')):
            return ImageProcessor()
        elif filename.endswith(('.xlsx', '.xls')):
            return ExcelProcessor()
        elif filename.endswith('.pdf'):
            return PDFProcessor()
        elif filename.endswith('.docx'):
            return DocxProcessor()
        elif filename.endswith(('.txt', '.md')):
            return TextProcessor()
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {filename}. Supported formats: PDF, TXT, DOC, DOCX, MD, XLSX, XLS, JPG, JPEG, PNG"
            )
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        """Process the file and extract text content"""
        raise NotImplementedError


class TextProcessor(DocumentProcessor):
    """Processor for text and markdown files"""
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        try:
            content = await file.read()
            text = content.decode('utf-8')
            
            return {
                'text': text,
                'metadata': {
                    'file_type': 'text',
                    'encoding': 'utf-8',
                    'length': len(text)
                }
            }
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail="Unable to decode text file as UTF-8")


class PDFProcessor(DocumentProcessor):
    """Processor for PDF files"""
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        try:
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            text = '\n'.join(text_parts)
            
            return {
                'text': text,
                'metadata': {
                    'file_type': 'pdf',
                    'pages': len(pdf_reader.pages),
                    'length': len(text)
                }
            }
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing PDF file: {str(e)}")


class DocxProcessor(DocumentProcessor):
    """Processor for DOCX files"""
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        try:
            content = await file.read()
            doc = DocxDocument(io.BytesIO(content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            text = '\n'.join(text_parts)
            
            return {
                'text': text,
                'metadata': {
                    'file_type': 'docx',
                    'paragraphs': len(doc.paragraphs),
                    'length': len(text)
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing DOCX file: {str(e)}")


class ExcelProcessor(DocumentProcessor):
    """Processor for Excel files (XLSX, XLS)"""
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        try:
            content = await file.read()
            workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            
            text_parts = []
            sheet_info = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                sheet_text = []
                row_count = 0
                
                # Add sheet name as header
                sheet_text.append(f"=== Sheet: {sheet_name} ===")
                
                for row in worksheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):  # Skip empty rows
                        row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                        sheet_text.append(row_text)
                        row_count += 1
                
                sheet_content = '\n'.join(sheet_text)
                text_parts.append(sheet_content)
                
                sheet_info.append({
                    'name': sheet_name,
                    'rows': row_count,
                    'columns': worksheet.max_column
                })
            
            text = '\n\n'.join(text_parts)
            
            # Flatten sheet info for ChromaDB compatibility
            sheet_names = ', '.join([info['name'] for info in sheet_info])
            total_rows = sum([info['rows'] for info in sheet_info])
            
            return {
                'text': text,
                'metadata': {
                    'file_type': 'excel',
                    'sheet_names': sheet_names,
                    'total_sheets': len(workbook.sheetnames),
                    'total_rows': total_rows,
                    'length': len(text)
                }
            }
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing Excel file: {str(e)}")


class ImageProcessor(DocumentProcessor):
    """Processor for image files using OCR"""
    
    async def process(self, file: UploadFile) -> Dict[str, Any]:
        try:
            content = await file.read()
            
            # Open image with Pillow
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use pytesseract to extract text
            text = pytesseract.image_to_string(image)
            
            # Clean up extracted text
            text = text.strip()
            if not text:
                text = "[No text detected in image]"
            
            return {
                'text': text,
                'metadata': {
                    'file_type': 'image',
                    'image_format': image.format or 'unknown',
                    'image_width': image.size[0],
                    'image_height': image.size[1],
                    'image_mode': image.mode,
                    'length': len(text),
                    'ocr_confidence': 'basic'  # pytesseract basic extraction
                }
            }
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing image file: {str(e)}. Make sure tesseract is installed on the system.")


# Convenience function for getting supported file types
def get_supported_extensions() -> Dict[str, list]:
    """Return dictionary of supported file extensions by category"""
    return {
        'text': ['.txt', '.md'],
        'documents': ['.pdf', '.docx'],
        'spreadsheets': ['.xlsx', '.xls'], 
        'images': ['.jpg', '.jpeg', '.png']
    }


def get_all_supported_extensions() -> list:
    """Return flat list of all supported extensions"""
    extensions = []
    for category in get_supported_extensions().values():
        extensions.extend(category)
    return extensions
